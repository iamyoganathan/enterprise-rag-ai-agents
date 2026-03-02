"""
Document Loader Module
Handles loading of various document formats (PDF, DOCX, TXT, Markdown).
"""

from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from abc import ABC, abstractmethod

from ..utils.logger import get_logger
from ..utils.config import get_settings

logger = get_logger(__name__)
settings = get_settings()


class BaseDocumentLoader(ABC):
    """Abstract base class for document loaders."""
    
    def __init__(self):
        """Initialize the document loader."""
        self.supported_extensions = []
    
    @abstractmethod
    def load(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Load a document from the given file path.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing document content and metadata
        """
        pass
    
    def supports(self, file_path: Union[str, Path]) -> bool:
        """
        Check if this loader supports the given file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if supported, False otherwise
        """
        extension = Path(file_path).suffix.lower()
        return extension in self.supported_extensions


class PDFLoader(BaseDocumentLoader):
    """Loader for PDF documents."""
    
    def __init__(self):
        """Initialize the PDF loader."""
        super().__init__()
        self.supported_extensions = ['.pdf']
    
    def load(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Load a PDF document.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary with content and metadata
        """
        try:
            import pymupdf  # PyMuPDF (fitz)
            
            file_path = Path(file_path)
            logger.info(f"Loading PDF: {file_path.name}")
            
            doc = pymupdf.open(str(file_path))
            
            # Extract text from all pages
            full_text = ""
            pages_content = []
            
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text() or ''
                full_text += text + "\n\n"
                pages_content.append({
                    "page_number": page_num,
                    "content": text,
                    "char_count": len(text)
                })
            
            # Extract metadata
            metadata = {
                "format": "pdf",
                "num_pages": len(doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
                "creator": doc.metadata.get("creator", ""),
                "producer": doc.metadata.get("producer", ""),
                "creation_date": doc.metadata.get("creationDate", ""),
                "modification_date": doc.metadata.get("modDate", ""),
            }
            
            doc.close()
            
            logger.info(f"Successfully loaded PDF: {file_path.name} ({metadata['num_pages']} pages)")
            
            return {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "content": full_text.strip(),
                "pages": pages_content,
                "metadata": metadata,
                "char_count": len(full_text),
                "word_count": len(full_text.split())
            }
            
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise


class DOCXLoader(BaseDocumentLoader):
    """Loader for DOCX documents."""
    
    def __init__(self):
        """Initialize the DOCX loader."""
        super().__init__()
        self.supported_extensions = ['.docx', '.doc']
    
    def load(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Load a DOCX document.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary with content and metadata
        """
        try:
            from docx import Document
            
            file_path = Path(file_path)
            logger.info(f"Loading DOCX: {file_path.name}")
            
            doc = Document(str(file_path))
            
            # Extract text from paragraphs
            paragraphs = []
            full_text = ""
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
                    full_text += text + "\n\n"
            
            # Extract text from tables
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_content.append(table_data)
                
                # Add table content to full text
                for row in table_data:
                    full_text += " | ".join(row) + "\n"
                full_text += "\n"
            
            # Extract metadata from core properties
            core_props = doc.core_properties
            metadata = {
                "format": "docx",
                "num_paragraphs": len(paragraphs),
                "num_tables": len(tables_content),
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            }
            
            logger.info(f"Successfully loaded DOCX: {file_path.name} ({len(paragraphs)} paragraphs)")
            
            return {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "content": full_text.strip(),
                "paragraphs": paragraphs,
                "tables": tables_content,
                "metadata": metadata,
                "char_count": len(full_text),
                "word_count": len(full_text.split())
            }
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise


class TextLoader(BaseDocumentLoader):
    """Loader for plain text documents."""
    
    def __init__(self):
        """Initialize the text loader."""
        super().__init__()
        self.supported_extensions = ['.txt', '.text']
    
    def load(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Load a text document.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Dictionary with content and metadata
        """
        try:
            file_path = Path(file_path)
            logger.info(f"Loading text file: {file_path.name}")
            
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError(f"Could not decode file with any of: {encodings}")
            
            # Count lines
            lines = content.split('\n')
            non_empty_lines = [line for line in lines if line.strip()]
            
            metadata = {
                "format": "txt",
                "encoding": used_encoding,
                "num_lines": len(lines),
                "num_non_empty_lines": len(non_empty_lines),
                "file_size_bytes": file_path.stat().st_size,
            }
            
            logger.info(f"Successfully loaded text file: {file_path.name} ({len(lines)} lines)")
            
            return {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "content": content,
                "lines": lines,
                "metadata": metadata,
                "char_count": len(content),
                "word_count": len(content.split())
            }
            
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            raise


class MarkdownLoader(BaseDocumentLoader):
    """Loader for Markdown documents."""
    
    def __init__(self):
        """Initialize the Markdown loader."""
        super().__init__()
        self.supported_extensions = ['.md', '.markdown']
    
    def load(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Load a Markdown document.
        
        Args:
            file_path: Path to the Markdown file
            
        Returns:
            Dictionary with content and metadata
        """
        try:
            import markdown
            from bs4 import BeautifulSoup
            
            file_path = Path(file_path)
            logger.info(f"Loading Markdown file: {file_path.name}")
            
            # Read raw markdown
            with open(file_path, 'r', encoding='utf-8') as f:
                raw_content = f.read()
            
            # Convert to HTML for structured parsing
            html = markdown.markdown(raw_content, extensions=['tables', 'fenced_code'])
            soup = BeautifulSoup(html, 'html.parser')
            
            # Extract plain text
            plain_text = soup.get_text()
            
            # Extract headers
            headers = []
            for i in range(1, 7):
                for header in soup.find_all(f'h{i}'):
                    headers.append({
                        "level": i,
                        "text": header.get_text()
                    })
            
            # Count code blocks
            code_blocks = len(soup.find_all('code'))
            
            metadata = {
                "format": "markdown",
                "num_headers": len(headers),
                "num_code_blocks": code_blocks,
                "has_tables": bool(soup.find('table')),
            }
            
            logger.info(f"Successfully loaded Markdown: {file_path.name} ({len(headers)} headers)")
            
            return {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "content": plain_text,
                "raw_markdown": raw_content,
                "html": html,
                "headers": headers,
                "metadata": metadata,
                "char_count": len(plain_text),
                "word_count": len(plain_text.split())
            }
            
        except Exception as e:
            logger.error(f"Error loading Markdown file {file_path}: {e}")
            raise


class DocumentLoaderFactory:
    """Factory for creating appropriate document loaders."""
    
    def __init__(self):
        """Initialize the document loader factory."""
        self.loaders = [
            PDFLoader(),
            DOCXLoader(),
            TextLoader(),
            MarkdownLoader()
        ]
    
    def get_loader(self, file_path: Union[str, Path]) -> Optional[BaseDocumentLoader]:
        """
        Get the appropriate loader for the given file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Appropriate document loader or None
        """
        for loader in self.loaders:
            if loader.supports(file_path):
                return loader
        return None
    
    def load_document(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Load a document using the appropriate loader.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Loaded document data
            
        Raises:
            ValueError: If no loader supports the file type
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        loader = self.get_loader(file_path)
        if loader is None:
            raise ValueError(f"No loader available for file type: {file_path.suffix}")
        
        return loader.load(file_path)
    
    def load_directory(
        self,
        directory: Union[str, Path],
        recursive: bool = True
    ) -> List[Dict[str, Any]]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory: Path to the directory
            recursive: Whether to search subdirectories
            
        Returns:
            List of loaded documents
        """
        directory = Path(directory)
        
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        if not directory.is_dir():
            raise ValueError(f"Not a directory: {directory}")
        
        documents = []
        pattern = "**/*" if recursive else "*"
        
        for file_path in directory.glob(pattern):
            if file_path.is_file() and self.get_loader(file_path):
                try:
                    doc = self.load_document(file_path)
                    documents.append(doc)
                except Exception as e:
                    logger.error(f"Failed to load {file_path}: {e}")
                    continue
        
        logger.info(f"Loaded {len(documents)} documents from {directory}")
        return documents


# Convenience function
def load_document(file_path: Union[str, Path]) -> Dict[str, Any]:
    """
    Load a document from the given file path.
    
    Args:
        file_path: Path to the document
        
    Returns:
        Loaded document data
    """
    factory = DocumentLoaderFactory()
    return factory.load_document(file_path)


def load_documents(directory: Union[str, Path], recursive: bool = True) -> List[Dict[str, Any]]:
    """
    Load all documents from a directory.
    
    Args:
        directory: Path to the directory
        recursive: Whether to search subdirectories
        
    Returns:
        List of loaded documents
    """
    factory = DocumentLoaderFactory()
    return factory.load_directory(directory, recursive)


if __name__ == "__main__":
    # Test the document loaders
    factory = DocumentLoaderFactory()
    
    # Test with sample directory
    sample_dir = Path("data/sample_documents")
    if sample_dir.exists():
        docs = factory.load_directory(sample_dir)
        print(f"Loaded {len(docs)} documents")
        
        for doc in docs:
            print(f"\n{doc['file_name']}:")
            print(f"  Format: {doc['metadata']['format']}")
            print(f"  Words: {doc['word_count']}")
            print(f"  First 100 chars: {doc['content'][:100]}...")
    else:
        print(f"Sample directory not found: {sample_dir}")
        print("Create the directory and add some documents to test.")
